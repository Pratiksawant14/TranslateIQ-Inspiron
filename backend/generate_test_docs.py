"""
Generate two large DOCX test documents for TranslateIQ pipeline testing:

Document A: "CloudSync Pro - Complete User Manual v1" (~300 segments)
  - Rich enterprise software documentation
  - Covers installation, configuration, features, troubleshooting

Document B: "CloudSync Pro - Administrator Guide v1" (~250 segments)
  - Related document with ~65% overlapping terminology/sentences from Doc A
  - Tests how well TM + local model reuse previous translations
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

OUTPUT_DIR = os.path.join(os.path.dirname(__file__), "..", "test_documents")
os.makedirs(OUTPUT_DIR, exist_ok=True)


def heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def para(doc, text):
    doc.add_paragraph(text)


def bullet(doc, items):
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")


def table_2col(doc, rows):
    table = doc.add_table(rows=len(rows)+1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Parameter"
    hdr[1].text = "Value"
    for i, (k, v) in enumerate(rows):
        table.rows[i+1].cells[0].text = k
        table.rows[i+1].cells[1].text = v


# ─── DOCUMENT A: Complete User Manual ────────────────────────────────────────

def create_document_a():
    doc = Document()
    doc.add_heading("CloudSync Pro — Complete User Manual v1.0", 0)

    heading(doc, "Chapter 1: Introduction", 1)
    para(doc, "CloudSync Pro is an enterprise-grade cloud synchronization and collaboration platform designed for organizations of all sizes. This manual covers installation, configuration, and day-to-day operation of CloudSync Pro version 1.0.")
    para(doc, "CloudSync Pro enables real-time file synchronization across all devices in your organization. The platform supports Windows, macOS, and Linux operating systems and integrates seamlessly with existing enterprise infrastructure.")
    para(doc, "Before proceeding with installation, please ensure your system meets the minimum hardware and software requirements described in this chapter.")

    heading(doc, "1.1 System Requirements", 2)
    para(doc, "CloudSync Pro requires a minimum system configuration to ensure stable and efficient operation. Systems that do not meet these requirements may experience degraded performance or installation failures.")
    para(doc, "RAM: CloudSync Pro requires at least 4GB of RAM to run properly. We recommend 8GB of RAM for best performance in enterprise environments.")
    para(doc, "Processor: A 64-bit processor running at 2GHz or faster is required. Multi-core processors are strongly recommended for optimal synchronization performance.")
    para(doc, "Storage: You need at least 500MB of free disk space for the CloudSync Pro installation files. Additional storage is required for your synchronized data.")
    para(doc, "Network: A stable broadband internet connection is required. We recommend a connection speed of at least 10 Mbps for efficient cloud synchronization.")
    para(doc, "Operating System: CloudSync Pro supports Windows 10, Windows 11, macOS 11 (Big Sur) and later, and Ubuntu 20.04 LTS and later distributions.")

    heading(doc, "1.2 Supported File Types", 2)
    para(doc, "CloudSync Pro supports synchronization of all common file types including documents, spreadsheets, presentations, images, audio files, and video files.")
    para(doc, "The following file types are supported for real-time collaboration: Microsoft Office documents (.docx, .xlsx, .pptx), PDF files, plain text files, and CSV files.")
    para(doc, "Binary files including executables, compressed archives, and encrypted files are synchronized but cannot be opened for real-time collaboration.")

    heading(doc, "Chapter 2: Installation Guide", 1)
    para(doc, "This chapter guides you through the complete installation process for CloudSync Pro. Follow each step carefully to ensure a successful installation.")
    para(doc, "Before beginning installation, close all running applications and save any open documents. The installation process may require a system restart to complete.")

    heading(doc, "2.1 Downloading CloudSync Pro", 2)
    para(doc, "Visit the official CloudSync Pro website at www.cloudsync-pro.com to download the latest version of the installation package for your operating system.")
    para(doc, "After downloading, verify the integrity of the installation package using the SHA-256 checksum provided on the download page. This ensures the file has not been corrupted during download.")
    para(doc, "The installation package for Windows is approximately 250MB. The macOS package is approximately 280MB. The Linux package is approximately 200MB.")

    heading(doc, "2.2 Windows Installation", 2)
    para(doc, "Double-click the downloaded CloudSync-Pro-Setup.exe file to launch the installation wizard. If prompted by User Account Control, click Yes to allow the installer to make changes to your system.")
    para(doc, "The installation wizard will guide you through a series of screens. On the License Agreement screen, read the terms carefully and click Accept to continue.")
    para(doc, "Choose your installation directory. The default installation path is C:\\Program Files\\CloudSync Pro. We recommend keeping the default path unless you have specific requirements.")
    para(doc, "Select the components you wish to install. The Core Synchronization Engine is required. Additional optional components include the Desktop Integration Extension and the Browser Plugin.")
    para(doc, "Click Install to begin the installation process. The installation typically takes between 5 and 15 minutes depending on your system performance.")
    para(doc, "When the installation is complete, you will see a confirmation screen. Click Finish to close the installer. A system restart may be required.")

    heading(doc, "2.3 macOS Installation", 2)
    para(doc, "Open the downloaded CloudSync-Pro.dmg file. Drag the CloudSync Pro application icon to the Applications folder to install it.")
    para(doc, "The first time you launch CloudSync Pro, macOS may ask you to confirm that you want to open an application downloaded from the internet. Click Open to proceed.")
    para(doc, "CloudSync Pro requires full disk access to synchronize files from all locations on your Mac. Grant this permission in System Preferences under Security and Privacy.")

    heading(doc, "2.4 Linux Installation", 2)
    para(doc, "CloudSync Pro is available as a .deb package for Debian-based distributions and as a .rpm package for Red Hat-based distributions.")
    para(doc, "For Ubuntu and Debian systems, open a terminal and run the following command: sudo dpkg -i cloudsync-pro_1.0_amd64.deb")
    para(doc, "After installation, start the CloudSync Pro service by running: sudo systemctl start cloudsync-pro")
    para(doc, "To enable CloudSync Pro to start automatically when your system boots, run: sudo systemctl enable cloudsync-pro")

    heading(doc, "Chapter 3: Initial Configuration", 1)
    para(doc, "After successful installation, you must configure CloudSync Pro before you can begin synchronizing files. The initial configuration wizard will guide you through this process.")
    para(doc, "Open CloudSync Pro from your desktop shortcut or application menu. The first time you open the application, the Initial Configuration Wizard will launch automatically.")

    heading(doc, "3.1 Creating Your Account", 2)
    para(doc, "If you do not have a CloudSync Pro account, click Create New Account on the welcome screen. You will need to provide your name, company email address, and a secure password.")
    para(doc, "Your password must be at least 12 characters long and must include uppercase letters, lowercase letters, numbers, and special characters. This is required for security compliance.")
    para(doc, "After creating your account, you will receive a confirmation email at the address you provided. Click the verification link in the email to activate your account.")
    para(doc, "If your organization uses Single Sign-On (SSO), contact your system administrator to configure SSO integration before attempting to log in.")

    heading(doc, "3.2 Connecting to Your Organization", 2)
    para(doc, "After logging in, CloudSync Pro will automatically detect your organization based on your email domain. If automatic detection fails, enter your organization ID manually.")
    para(doc, "Your system administrator will provide you with your organization ID and the server address. Enter these details in the Organization Settings screen.")
    para(doc, "CloudSync Pro uses end-to-end encryption for all data transfers. Your encryption keys are generated during this initial setup and are stored securely on your device.")

    heading(doc, "3.3 Choosing Sync Folders", 2)
    para(doc, "Select the folders on your device that you want to synchronize with CloudSync Pro. You can select multiple folders from different locations on your system.")
    para(doc, "The Desktop, Documents, and Downloads folders are pre-selected by default. You can deselect them and choose custom folders according to your requirements.")
    para(doc, "Be careful when selecting folders for synchronization. All files in selected folders will be uploaded to the cloud and shared according to your organization's permissions settings.")

    heading(doc, "Chapter 4: Core Features", 1)
    para(doc, "CloudSync Pro provides a comprehensive suite of features for enterprise file management and collaboration. This chapter describes all core features in detail.")

    heading(doc, "4.1 Real-Time Synchronization", 2)
    para(doc, "CloudSync Pro synchronizes your files in real time as soon as changes are detected. Changes made on any device are reflected on all other devices within seconds.")
    para(doc, "The synchronization engine uses delta sync technology, which means only the changed portions of files are transferred rather than the entire file. This saves bandwidth and speeds up synchronization.")
    para(doc, "You can monitor the real-time synchronization status in the system tray icon. A spinning icon indicates synchronization is in progress. A checkmark icon indicates all files are up to date.")
    para(doc, "CloudSync Pro maintains a local copy of all synchronized files so you can access them even when you do not have an internet connection.")

    heading(doc, "4.2 Version History", 2)
    para(doc, "CloudSync Pro automatically maintains a version history for all synchronized files. You can access the version history by right-clicking any file and selecting Version History.")
    para(doc, "By default, CloudSync Pro keeps the last 50 versions of each file. Your administrator can configure this limit based on your organization's storage plan.")
    para(doc, "To restore a previous version, select the version you want from the Version History panel and click Restore. The selected version will replace the current version of the file.")
    para(doc, "You can also download a specific version without replacing the current version by selecting the version and clicking Download.")

    heading(doc, "4.3 File Sharing and Collaboration", 2)
    para(doc, "CloudSync Pro allows you to share files and folders with colleagues inside your organization and with external partners outside your organization.")
    para(doc, "To share a file or folder, right-click it and select Share. The Share dialog allows you to specify recipients and set permissions such as view-only or edit access.")
    para(doc, "Shared folders appear in the Shared with Me section of the CloudSync Pro application. Changes made by any collaborator are visible to all other collaborators in real time.")
    para(doc, "When sharing with external users, CloudSync Pro generates a secure sharing link. You can set an expiration date and a password for external sharing links.")

    heading(doc, "4.4 Offline Access", 2)
    para(doc, "CloudSync Pro stores a local copy of your synchronized files on your device, allowing you to access and edit them without an internet connection.")
    para(doc, "Files modified while offline are queued for synchronization and automatically uploaded when your internet connection is restored. Conflict detection alerts you if the same file was modified by another user while you were offline.")
    para(doc, "You can manually mark specific files or folders as Available Offline to ensure they are always stored locally on your device.")

    heading(doc, "Chapter 5: Security and Compliance", 1)
    para(doc, "Security is a core principle of CloudSync Pro. The platform is designed to meet the stringent security and compliance requirements of enterprise organizations.")

    heading(doc, "5.1 Encryption", 2)
    para(doc, "All data stored in CloudSync Pro is encrypted using AES-256 encryption, which is the industry standard for protecting sensitive information.")
    para(doc, "Data in transit between your device and CloudSync Pro servers is protected using TLS 1.3. This ensures that data cannot be intercepted during transmission.")
    para(doc, "Encryption keys are managed per organization and are never shared with CloudSync Pro Inc. Your data remains private and accessible only to authorized users within your organization.")

    heading(doc, "5.2 Access Controls", 2)
    para(doc, "CloudSync Pro provides role-based access control (RBAC) that allows administrators to define granular permissions for users and groups.")
    para(doc, "Administrators can assign the following roles: Owner, Administrator, Editor, Contributor, and Viewer. Each role has a specific set of permissions.")
    para(doc, "Access can be revoked instantly for any user. When access is revoked, the user's local copy of shared files is deleted automatically from their device.")
    para(doc, "Multi-factor authentication is supported and strongly recommended for all user accounts. Administrators can make MFA mandatory for all users in their organization.")

    heading(doc, "5.3 Audit Logs", 2)
    para(doc, "CloudSync Pro maintains detailed audit logs of all user activities. Administrators can access these logs to review file access, modifications, sharing events, and login activities.")
    para(doc, "Audit logs are retained for 12 months by default. Organizations with compliance requirements can request extended log retention of up to 7 years.")
    para(doc, "Audit logs can be exported in CSV format for integration with third-party Security Information and Event Management (SIEM) systems.")

    heading(doc, "Chapter 6: Troubleshooting", 1)
    para(doc, "This chapter provides solutions for the most common issues that users encounter when using CloudSync Pro.")

    heading(doc, "6.1 Synchronization Issues", 2)
    para(doc, "If files are not synchronizing correctly, first check that your device is connected to the internet and that the CloudSync Pro service is running.")
    para(doc, "Open Task Manager on Windows or Activity Monitor on macOS to verify that the CloudSync Pro process is running. If it is not running, restart the application.")
    para(doc, "If synchronization is stuck and not progressing, try pausing and resuming synchronization from the CloudSync Pro menu in the system tray.")
    para(doc, "For persistent synchronization issues, check the application logs located in the CloudSync Pro installation directory under the Logs subfolder.")

    heading(doc, "6.2 Login and Authentication Issues", 2)
    para(doc, "If you cannot log in to CloudSync Pro, verify that your email address and password are correct. Passwords are case-sensitive.")
    para(doc, "If you have forgotten your password, click Forgot Password on the login screen. You will receive a password reset email within 5 minutes.")
    para(doc, "If your account has been locked after too many failed login attempts, contact your system administrator or CloudSync Pro support to unlock it.")

    heading(doc, "6.3 Performance Issues", 2)
    para(doc, "If CloudSync Pro is consuming too much CPU or memory, you can limit its resource usage in the Settings menu under Performance.")
    para(doc, "Reducing the number of files in your synchronized folders and excluding large files from synchronization can significantly improve performance.")
    para(doc, "If the application is slow to start, disable any startup items that are not required by going to Settings and then Startup Preferences.")

    heading(doc, "Chapter 7: Advanced Configuration", 1)
    para(doc, "This chapter describes advanced configuration options for power users and system administrators.")

    heading(doc, "7.1 Bandwidth Management", 2)
    para(doc, "CloudSync Pro allows you to set upload and download bandwidth limits to prevent the synchronization engine from consuming your entire internet connection.")
    para(doc, "You can configure different bandwidth limits for different times of day using the Bandwidth Scheduler in the Settings menu.")
    para(doc, "Limiting upload bandwidth to 50% of your available connection speed is recommended to ensure other applications continue to function normally during synchronization.")

    heading(doc, "7.2 Proxy Configuration", 2)
    para(doc, "If your organization uses a network proxy, configure the proxy settings in CloudSync Pro under Settings and then Network.")
    para(doc, "CloudSync Pro supports HTTP, HTTPS, and SOCKS5 proxy configurations. Enter the proxy host address, port number, and authentication credentials if required.")
    para(doc, "Contact your network administrator if you are unsure about your organization's proxy settings.")

    heading(doc, "7.3 Command Line Interface", 2)
    para(doc, "CloudSync Pro provides a command line interface (CLI) for advanced users and system administrators who prefer to manage synchronization from the terminal.")
    para(doc, "The CLI is available on all supported operating systems. On Windows, open Command Prompt and navigate to the CloudSync Pro installation directory to access the CLI.")
    para(doc, "Common CLI commands include: csync status to check synchronization status, csync pause to pause synchronization, and csync resume to resume synchronization.")
    para(doc, "Run csync help to see a complete list of available CLI commands and their options.")

    heading(doc, "Chapter 8: Support and Contact", 1)
    para(doc, "CloudSync Pro provides multiple support channels to help you resolve issues and get the most out of the platform.")
    para(doc, "Email support is available 24 hours a day, 7 days a week. Send your support request to support@cloudsync-pro.com with a detailed description of your issue.")
    para(doc, "Phone support is available Monday through Friday from 9:00 AM to 6:00 PM Eastern Time. Call our support line at 1-800-CLOUDSYNC.")
    para(doc, "The CloudSync Pro documentation portal at docs.cloudsync-pro.com contains comprehensive guides, tutorials, and frequently asked questions.")
    para(doc, "Enterprise customers have access to a dedicated support engineer who can provide personalized assistance and proactive monitoring of your CloudSync Pro deployment.")

    path = os.path.join(OUTPUT_DIR, "CloudSync_Pro_User_Manual_v1.docx")
    doc.save(path)
    print(f"Document A saved: {path}")
    return path


# ─── DOCUMENT B: Administrator Guide (overlapping vocabulary from Doc A) ──────

def create_document_b():
    doc = Document()
    doc.add_heading("CloudSync Pro — Administrator Guide v1.0", 0)

    heading(doc, "Chapter 1: Administrator Overview", 1)
    para(doc, "This Administrator Guide provides detailed instructions for IT administrators responsible for deploying, configuring, and maintaining CloudSync Pro within an enterprise environment.")
    para(doc, "CloudSync Pro is an enterprise-grade cloud synchronization and collaboration platform designed for organizations of all sizes. Administrators play a critical role in ensuring the platform runs efficiently and securely.")
    para(doc, "This guide covers user management, security configuration, server administration, monitoring, and compliance reporting.")

    heading(doc, "1.1 Administrative Prerequisites", 2)
    para(doc, "To administer CloudSync Pro, you must have administrator privileges in your organization's CloudSync Pro account. Contact your CloudSync Pro sales representative to request administrative access.")
    para(doc, "RAM: CloudSync Pro requires at least 4GB of RAM to run properly. For the server components, we recommend 16GB of RAM for enterprise deployments handling more than 100 concurrent users.")
    para(doc, "Network: A stable broadband internet connection is required for the administrator console. We recommend a dedicated management network segment for administrative access.")
    para(doc, "Operating System: The CloudSync Pro administrator console supports Windows 10, Windows 11, macOS 11 (Big Sur) and later.")

    heading(doc, "1.2 Administrative Console Overview", 2)
    para(doc, "The CloudSync Pro Administrator Console provides a web-based interface for managing all aspects of your organization's CloudSync Pro deployment.")
    para(doc, "Access the Administrator Console by navigating to admin.cloudsync-pro.com and logging in with your administrator credentials.")
    para(doc, "The console is organized into the following sections: Dashboard, User Management, Security, Sync Policies, Storage, Compliance, and Support.")

    heading(doc, "Chapter 2: User Management", 1)
    para(doc, "Effective user management is essential for maintaining security and controlling access to your organization's data in CloudSync Pro.")

    heading(doc, "2.1 Adding Users", 2)
    para(doc, "To add a new user, navigate to User Management in the Administrator Console and click Add User.")
    para(doc, "Enter the user's full name, company email address, and assign the appropriate role. Available roles are: Owner, Administrator, Editor, Contributor, and Viewer.")
    para(doc, "If your organization uses Single Sign-On (SSO), new users can be provisioned automatically through your identity provider. Contact your system administrator to configure SSO integration.")
    para(doc, "Your password must be at least 12 characters long and must include uppercase letters, lowercase letters, numbers, and special characters. This is required for security compliance.")

    heading(doc, "2.2 Group Management", 2)
    para(doc, "Groups allow you to manage permissions for multiple users simultaneously. Create groups based on departments, projects, or any organizational structure that suits your needs.")
    para(doc, "To create a group, go to User Management, select Groups, and click Create Group. Enter a name and description for the group.")
    para(doc, "Add users to groups by selecting the group and clicking Add Members. You can add individual users or import a list of users from a CSV file.")
    para(doc, "Group permissions are additive. A user who belongs to multiple groups inherits the combined permissions of all groups they belong to.")

    heading(doc, "2.3 Removing Users", 2)
    para(doc, "When an employee leaves your organization, immediately revoke their CloudSync Pro access. Go to User Management, find the user, and click Deactivate Account.")
    para(doc, "Access can be revoked instantly for any user. When access is revoked, the user's local copy of shared files is deleted automatically from their device.")
    para(doc, "Deactivated users cannot log in to CloudSync Pro and cannot access any files or folders. Their files remain in the system and are reassigned to their manager by default.")

    heading(doc, "Chapter 3: Security Administration", 1)
    para(doc, "Security is a core principle of CloudSync Pro. As an administrator, you are responsible for configuring and maintaining the security settings of your organization's deployment.")

    heading(doc, "3.1 Encryption Management", 2)
    para(doc, "All data stored in CloudSync Pro is encrypted using AES-256 encryption, which is the industry standard for protecting sensitive information.")
    para(doc, "Data in transit between your device and CloudSync Pro servers is protected using TLS 1.3. This ensures that data cannot be intercepted during transmission.")
    para(doc, "As an administrator, you can view and rotate your organization's encryption keys in the Security section of the Administrator Console.")
    para(doc, "Encryption keys are managed per organization and are never shared with CloudSync Pro Inc. Your data remains private and accessible only to authorized users within your organization.")

    heading(doc, "3.2 Configuring Multi-Factor Authentication", 2)
    para(doc, "Multi-factor authentication is supported and strongly recommended for all user accounts. Administrators can make MFA mandatory for all users in their organization.")
    para(doc, "To enable mandatory MFA, go to Security in the Administrator Console and toggle the Require Multi-Factor Authentication setting to On.")
    para(doc, "CloudSync Pro supports authenticator apps, SMS verification, and hardware security keys as MFA methods. We recommend authenticator apps for the best balance of security and usability.")
    para(doc, "Users who have not set up MFA will be prompted to do so on their next login after MFA is made mandatory.")

    heading(doc, "3.3 Managing Audit Logs", 2)
    para(doc, "CloudSync Pro maintains detailed audit logs of all user activities. Administrators can access these logs to review file access, modifications, sharing events, and login activities.")
    para(doc, "Access the Audit Logs by navigating to Compliance in the Administrator Console. You can filter logs by user, date range, action type, and file or folder.")
    para(doc, "Audit logs are retained for 12 months by default. Organizations with compliance requirements can request extended log retention of up to 7 years.")
    para(doc, "Audit logs can be exported in CSV format for integration with third-party Security Information and Event Management (SIEM) systems.")

    heading(doc, "Chapter 4: Synchronization Policies", 1)
    para(doc, "Synchronization policies allow you to control what files are synchronized, how synchronization behaves, and how much bandwidth it uses across your organization.")

    heading(doc, "4.1 File Type Policies", 2)
    para(doc, "Use file type policies to control which types of files can be synchronized. Restricting synchronization to specific file types reduces storage usage and potential security risks.")
    para(doc, "The following file types are supported for real-time collaboration: Microsoft Office documents, PDF files, plain text files, and CSV files.")
    para(doc, "You can block synchronization of executable files (.exe, .bat, .sh) to prevent potential malware distribution through the platform.")
    para(doc, "File size limits can be set on a per-group or per-user basis. The maximum file size for synchronization is 50GB per file.")

    heading(doc, "4.2 Bandwidth Policies", 2)
    para(doc, "CloudSync Pro allows you to set upload and download bandwidth limits to prevent the synchronization engine from consuming your entire internet connection.")
    para(doc, "Set organization-wide bandwidth limits in the Sync Policies section of the Administrator Console. Individual users can further restrict their personal limits within the organization maximum.")
    para(doc, "You can configure different bandwidth limits for different times of day using the Bandwidth Scheduler. Set lower limits during business hours and higher limits overnight.")
    para(doc, "Limiting upload bandwidth to 50% of your available connection speed is recommended to ensure other applications continue to function normally during synchronization.")

    heading(doc, "4.3 Retention Policies", 2)
    para(doc, "Version history retention policies determine how many versions of each file are stored and for how long.")
    para(doc, "By default, CloudSync Pro keeps the last 50 versions of each file. Your administrator can configure this limit based on your organization's storage plan.")
    para(doc, "Set a maximum retention period to automatically delete old versions after a specified number of days. This helps control storage costs.")

    heading(doc, "Chapter 5: Storage Management", 1)
    para(doc, "Effective storage management ensures that your organization's CloudSync Pro deployment remains cost-effective and performs well.")

    heading(doc, "5.1 Storage Quotas", 2)
    para(doc, "Assign storage quotas to users and groups to control storage usage. Navigate to User Management and select a user or group to set their storage quota.")
    para(doc, "When a user approaches their storage quota, CloudSync Pro will send them email notifications at 80% and 95% of their quota.")
    para(doc, "Users who have reached their storage quota cannot upload new files until they free up space by deleting files or until their quota is increased by an administrator.")

    heading(doc, "5.2 Storage Reports", 2)
    para(doc, "Access storage reports in the Compliance section of the Administrator Console to monitor storage usage across your organization.")
    para(doc, "Reports show total storage usage, storage usage per user, storage usage per department, and storage growth trends over time.")
    para(doc, "Export storage reports in CSV format for further analysis or to share with management.")

    heading(doc, "Chapter 6: Monitoring and Troubleshooting", 1)
    para(doc, "Regular monitoring of your CloudSync Pro deployment helps you identify and resolve issues before they impact users.")

    heading(doc, "6.1 System Health Dashboard", 2)
    para(doc, "The System Health Dashboard in the Administrator Console provides a real-time overview of your CloudSync Pro deployment's performance and status.")
    para(doc, "Key metrics displayed include active users, synchronization queue depth, error rates, and storage capacity.")
    para(doc, "Set up email alerts to be notified when specific metrics exceed threshold values. For example, you can receive an alert when the error rate exceeds 5%.")

    heading(doc, "6.2 Resolving Common User Issues", 2)
    para(doc, "If files are not synchronizing correctly, first check that the user's device is connected to the internet and that the CloudSync Pro service is running.")
    para(doc, "Open Task Manager on Windows or Activity Monitor on macOS to verify that the CloudSync Pro process is running. If it is not running, restart the application.")
    para(doc, "If synchronization is stuck and not progressing, try pausing and resuming synchronization from the CloudSync Pro menu in the system tray.")
    para(doc, "For persistent synchronization issues, check the application logs located in the CloudSync Pro installation directory under the Logs subfolder.")

    heading(doc, "6.3 Escalating Issues to Support", 2)
    para(doc, "Email support is available 24 hours a day, 7 days a week. Send your support request to support@cloudsync-pro.com with a detailed description of your issue.")
    para(doc, "Enterprise customers have access to a dedicated support engineer who can provide personalized assistance and proactive monitoring of your CloudSync Pro deployment.")
    para(doc, "When escalating issues, include the organization ID, affected user accounts, error messages, and logs from the CloudSync Pro Logs directory.")

    heading(doc, "Chapter 7: Compliance and Reporting", 1)
    para(doc, "CloudSync Pro provides comprehensive compliance features to help your organization meet regulatory requirements.")

    heading(doc, "7.1 GDPR Compliance", 2)
    para(doc, "CloudSync Pro is fully compliant with the General Data Protection Regulation (GDPR). All personal data is stored in EU data centers by default for European organizations.")
    para(doc, "Audit logs can be exported in CSV format for integration with third-party Security Information and Event Management (SIEM) systems.")
    para(doc, "Use the Data Subject Request tool in the Compliance section to handle GDPR data access and deletion requests from users.")

    heading(doc, "7.2 Compliance Reports", 2)
    para(doc, "Generate compliance reports directly from the Administrator Console. Reports include user activity summaries, access control changes, and data transfer logs.")
    para(doc, "Audit logs are retained for 12 months by default. Organizations with compliance requirements can request extended log retention of up to 7 years.")
    para(doc, "Schedule automatic report generation to ensure compliance reports are produced at regular intervals without manual intervention.")

    heading(doc, "Chapter 8: Disaster Recovery", 1)
    para(doc, "A robust disaster recovery plan is essential for maintaining business continuity in the event of a system failure or data loss.")

    heading(doc, "8.1 Backup Configuration", 2)
    para(doc, "CloudSync Pro automatically creates backups of all organizational data. Configure backup schedules in the Settings section of the Administrator Console.")
    para(doc, "By default, CloudSync Pro performs daily incremental backups and weekly full backups. These backups are stored in a separate geographic region from your primary data.")
    para(doc, "Verify that your backup configuration meets your organization's Recovery Point Objective (RPO) and Recovery Time Objective (RTO) requirements.")

    heading(doc, "8.2 Recovery Procedures", 2)
    para(doc, "To restore data from a backup, navigate to Compliance, then Recovery in the Administrator Console and select the backup point you want to restore from.")
    para(doc, "To restore a previous version, select the version you want from the Version History panel and click Restore. The selected version will replace the current version of the file.")
    para(doc, "Test your disaster recovery procedures at least twice per year to ensure that recovery objectives can be met and that all personnel are familiar with recovery processes.")

    path = os.path.join(OUTPUT_DIR, "CloudSync_Pro_Admin_Guide_v1.docx")
    doc.save(path)
    print(f"Document B saved: {path}")
    return path


if __name__ == "__main__":
    path_a = create_document_a()
    path_b = create_document_b()
    
    size_a = os.path.getsize(path_a) / 1024
    size_b = os.path.getsize(path_b) / 1024
    
    print(f"\nDocument A: {size_a:.1f} KB")
    print(f"Document B: {size_b:.1f} KB")
    print("\nBoth documents ready for upload to TranslateIQ!")
    print("Document A = Primary training data (upload + approve all)")
    print("Document B = Test document (should show high TM match rate after Doc A approval)")
