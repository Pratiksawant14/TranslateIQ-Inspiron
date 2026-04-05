import os
import csv
from docx import Document

def create_doc(filename, title, content_lines, list_items=None):
    doc = Document()
    doc.add_heading(title, 0)
    
    for line in content_lines:
        doc.add_paragraph(line)
        
    if list_items:
        for item in list_items:
            doc.add_paragraph(item, style='List Bullet')
            
    doc.save(filename)
    print(f"Created: {filename}")

def main():
    base_dir = r"c:\Users\Samadhan\Desktop\TranslateIQ-Inspiron\test_documents\presentation"
    os.makedirs(base_dir, exist_ok=True)
    
    # DOCUMENT 1: Baseline (100% New)
    doc1_lines = [
        "Simplify Healthcare Overview",
        "Simplify Healthcare is a technology solutions company that provides cloud-based software designed for the health insurance industry.",
        "Our mission is to replace manual processes with a single, integrated digital platform.",
        "We help Healthcare Payers automate complex back-office operations.",
        "We provide tools to digitally build and manage insurance benefit packages.",
        "This eliminates the need for slow, error-prone manual spreadsheets.",
        "Our document generation modules ensure compliance with CMS guidelines.",
        "The platform ensures benefits and provider data are accurately keyed into administrative systems.",
        "Simplify Healthcare helps health plans reduce operational costs and get new products to market faster."
    ]
    doc1_list = [
        "Benefit Plan Management",
        "Provider Lifecycle and Contract Management",
        "Claims Configuration",
        "Inquiry Management"
    ]
    create_doc(os.path.join(base_dir, "01_Simplify_Healthcare_Overview.docx"), "Simplify Healthcare Platform Overview", doc1_lines, doc1_list)

    # DOCUMENT 2: Evolution (Mix of Exact, Fuzzy, and New)
    doc2_lines = [
        "Simplify Healthcare Provider Operations Guide",  # Fuzzy title
        "Simplify Healthcare is a technology solutions company that provides cloud-based software designed for the health insurance industry.", # EXACT
        "Our primary focus is to replace manual provider credentialing with a single, integrated digital platform.", # FUZZY
        "We help Healthcare Payers automate complex provider data operations.", # FUZZY
        "Provider Lifecycle Management is a core business area.", # NEW
        "The platform automates processes such as creating network rosters and maintaining provider contracts.", # NEW
        "Our document generation modules ensure compliance with the No Surprises Act.", # FUZZY
        "This ensures robust regulatory compliance and enhances the overall member experience.", # NEW
        "Simplify Healthcare helps health plans reduce operational costs and get new products to market faster." # EXACT
    ]
    create_doc(os.path.join(base_dir, "02_Simplify_Healthcare_Provider_Guide.docx"), "Simplify Healthcare Provider Guide", doc2_lines)

    # DOCUMENT 3: High Matching retention (Mostly Exact/Fuzzy)
    doc3_lines = [
        "Simplify Healthcare Claims Addendum Guide v2", # FUZZY
        "Simplify Healthcare is a technology solutions company that provides AI-powered software designed for the health insurance industry.", # FUZZY (added AI-powered)
        "We help Healthcare Payers automate complex back-office operations.", # EXACT (From doc 1)
        "Claims Configuration ensures that benefits and provider data are accurately keyed into administrative systems.", # FUZZY 
        "This eliminates the need for slow, error-prone manual spreadsheets.", # EXACT (From doc 1)
        "Our inquiry management uses conversational AI to improve how providers look up benefits.", # NEW
        "This ensures robust regulatory compliance and enhances the overall member experience.", # EXACT (From doc 2)
        "Simplify Healthcare helps health plans reduce operational costs and get new products to market faster." # EXACT (From doc 1/2)
    ]
    create_doc(os.path.join(base_dir, "03_Simplify_Healthcare_Claims_v2.docx"), "Simplify Claims Addendum v2", doc3_lines)

    # CREATE GLOSSARY CSV
    glossary_path = os.path.join(base_dir, "simplify_glossary.csv")
    with open(glossary_path, mode='w', newline='', encoding='utf-8') as f:
        writer = csv.writer(f)
        writer.writerow(["source_term", "target_term", "context_notes"])
        writer.writerow(["Simplify Healthcare", "Simplify Healthcare", "Do not translate company name"])
        writer.writerow(["Healthcare Payers", "Pagadores de Atención Médica", "Industry term"])
        writer.writerow(["Benefit Plan Management", "Gestión de Planes de Beneficios", "Core product module"])
        writer.writerow(["Provider Lifecycle", "Ciclo de Vida del Proveedor", "Core product module"])
        writer.writerow(["Claims Configuration", "Configuración de Reclamos", "Core product module"])
        writer.writerow(["No Surprises Act", "Ley Sin Sorpresas", "US federal regulation"])
        writer.writerow(["CMS guidelines", "Pautas de CMS", "US federal regulation abbreviation"])
        writer.writerow(["conversational AI", "IA conversacional", "Tech terminology"])
        
    print(f"Created Glossary: {glossary_path}")

if __name__ == "__main__":
    main()
