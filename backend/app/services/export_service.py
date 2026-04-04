import os
from datetime import datetime
from fastapi import HTTPException, status
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select
from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor

from app.models.document import Document
from app.models.segment import Segment
from app.config import settings

def _get_final_text(segment: Segment) -> str:
    if segment.status in ["approved", "edited"] and segment.translated_text:
        return segment.translated_text
    return segment.source_text

async def export_document_docx(db: AsyncSession, document_id: str) -> str:
    doc_query = await db.execute(select(Document).where(Document.id == document_id))
    document = doc_query.scalars().first()
    
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
        
    segments_query = await db.execute(
        select(Segment)
        .where(Segment.document_id == document_id)
        .order_by(Segment.segment_index)
    )
    segments = segments_query.scalars().all()
    
    approved_count = sum(1 for s in segments if s.status in ["approved", "edited"])
    
    if approved_count == 0:
        raise HTTPException(status_code=400, detail="No approved segments to export.")
        
    doc = DocxDocument()
    
    # default font setup
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Inter'
    font.size = Pt(11)
    
    idx = 0
    while idx < len(segments):
        seg = segments[idx]
        text = _get_final_text(seg)
        ctype = seg.content_type
        
        if ctype == "heading":
            heading = doc.add_heading(text, level=1)
            for run in heading.runs:
                run.font.name = 'Inter'
                run.font.size = Pt(16)
                run.font.bold = True
                run.font.color.rgb = RGBColor(15, 27, 45)  # #0F1B2D
            idx += 1
            
        elif ctype == "list":
            doc.add_paragraph(text, style='List Bullet')
            idx += 1
            
        elif ctype == "table":
            # Collect consecutive table segments
            table_segments = [seg]
            next_idx = idx + 1
            while next_idx < len(segments) and segments[next_idx].content_type == "table":
                table_segments.append(segments[next_idx])
                next_idx += 1
                
            table = doc.add_table(rows=0, cols=1)
            table.style = 'Table Grid'
            
            for t_seg in table_segments:
                row_cells = table.add_row().cells
                row_cells[0].text = _get_final_text(t_seg)
                
            idx = next_idx
            
        else: # paragraph or anything else
            doc.add_paragraph(text)
            idx += 1
            
    # Add Export Metadata Footer
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0]
    
    now_str = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    footer_para.text = f"Exported by TranslateIQ\t\t{now_str}"
    
    # Save file
    os.makedirs(settings.EXPORTS_DIR, exist_ok=True)
    
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"{document_id}_translated_{timestamp}.docx"
    filepath = os.path.join(settings.EXPORTS_DIR, filename)
    
    doc.save(filepath)
    
    return filepath
