import os
try:
    import docx
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.oxml.ns import qn
except ImportError:
    print("Installing python-docx...")
    os.system("pip install python-docx")
    import docx
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def create_legal_contract():
    doc = docx.Document()
    
    # Title
    title = doc.add_heading('Legal Compliance Contract', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # 1. Introduction
    doc.add_heading('Introduction', level=1)
    doc.add_paragraph(
        'This Legal Compliance Contract (the "Agreement") is entered into by and between the Provider and the Client. '
        'This document serves as the official binding legal agreement governing the provision of professional services, software utilization, and associated support mechanisms. '
        'The primary purpose of this specific contract is to comprehensively outline the rules, regulations, and operational guidelines governing the use of the proprietary platform developed by the Provider. '
        'We strongly encourage all users and administrators to carefully review the complete documentation. '
        'Please refer to the User Guide for detailed programmatic instructions on effectively operating the software suite. '
        'It is highly expected and required that all parties comply fully with the technical procedures documented in the User Manual to ensure a seamless and error-free integration process. '
        'The platform is designed to be highly secure and fully compliant with international data privacy standards and security protocols.'
    )
    
    # 2. Terms and Conditions
    doc.add_heading('Terms and Conditions', level=1)
    doc.add_paragraph(
        'By executing and signing this comprehensive agreement, both involved parties openly acknowledge their respective duties and responsibilities limitlessly. '
        'The Client agrees to utilize the services strictly for lawful purposes and in accordance with all local regulations. '
        'There shall be absolutely no unauthorized  distribution, modification, or reverse engineering of this confidential document or the underlying proprietary technology. ' # Double space here (1)
        'Any such prohibited behavior or breach of these terms will immediately result in severe financial penalties and potential legal actions. '
        'For operational efficiency, we have established specific contact hours for our technical support division. '
        'All standard support inquiries and maintenance requests must be submitted by 5:00 PM on regular business days to guarantee a prompt response. '
        'However, we note that any service requests received after 5:00 P.M. will automatically be deferred and subsequently processed on the following business day. '
        'The Provider ensures high availability; nonetheless, scheduled maintenance windows are necessary. '
        'All operational obligations  under this extensive agreement remain in full and binding effect until the specific contract is formally and mutually concluded by both entities.' # Double space (2)
    )
    
    # 3. Payment Terms
    doc.add_heading('Payment Terms', level=1)
    doc.add_paragraph(
        'Financial compensation for the meticulously provided services must be settled accurately and strictly according to the customized milestone schedule. '
        'This agreement necessitates structured fiscal responsibility.'
    )
    
    # Table
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Milestone'
    hdr_cells[1].text = 'Amount'
    hdr_cells[2].text = 'Date'
    
    row_cells = table.rows[1].cells
    row_cells[0].text = 'Initial Project Kickoff'
    row_cells[1].text = '$1,000'
    row_cells[2].text = '11/15/2026' # MM/DD/YYYY
    
    row_cells = table.rows[2].cells
    row_cells[0].text = 'Midpoint Deliverable Review'
    row_cells[1].text = '$2,500'
    row_cells[2].text = '15-12-2026' # DD-MM-YYYY
    
    row_cells = table.rows[3].cells
    row_cells[0].text = 'Final Product Handover'
    row_cells[1].text = '$1,500'
    row_cells[2].text = '01/30/2027' # MM/DD/YYYY
    
    doc.add_paragraph() # Add some spacing
    doc.add_paragraph(
        'Failure to adequately make any required payments by the explicitly specified dates outlined in the table will automatically incur an compounding late fee of exactly five percent per standard calendar month until the pending balance is fully resolved. '
        'Both parties commit to transparency. All processed transactions  will be securely and permanently recorded in the official designated financial ledger for compliance auditing. ' # Double space (3)
        'Payment shall be made within thirty days.'
    )
    
    # 4. Obligations
    doc.add_heading('Obligations', level=1)
    doc.add_paragraph('To ensure mutual success, the following specific commitments apply:')
    doc.add_paragraph('The client must consistently and explicitly provide timely feedback during review cycles.', style='List Bullet')
    doc.add_paragraph('The provider must rigorously ensure and maintain a minimum of 99.9% server uptime.', style='List Bullet')
    doc.add_paragraph('Both parties will unequivocally maintain strict and perpetual confidentiality regarding all shared trade secrets.', style='List Bullet')
    doc.add_paragraph('Neither entity shall solicit or hire the employees of the other for a duration of two years.', style='List Bullet')
    doc.add_paragraph('The parties agree to the following terms and conditions.', style='List Bullet')
    
    # 5. Termination Clause
    doc.add_heading('Termination Clause', level=1)
    doc.add_paragraph(
        'Either party may terminate this agreement with written notice. '
        'In the specific event of an early or unexpected termination prompted by the Client, all outstanding and accumulated payments for completed work instantly become due immediately without exception. '
        'The Provider reserves the absolute right to suspend access if a material breach is not cured within ten days of notification.'
    )
    
    # 6. Governing Law
    doc.add_heading('Governing Law', level=1)
    doc.add_paragraph(
        'This agreement is governed by applicable law. '
        'Any potential disputes prominently arising from or directly relating to this contract will be exclusively and finally resolved in the designated state courts. '
        'Both parties formally waive their absolute right to a jury trial in any ensuing litigation.'
    )
    
    output_path = os.path.join(os.path.dirname(__file__), 'legal_contract.docx')
    doc.save(output_path)
    print(f"Successfully created {output_path}")

if __name__ == '__main__':
    create_legal_contract()
