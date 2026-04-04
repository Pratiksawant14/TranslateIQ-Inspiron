"""
STEP 3: Create test documents for full pipeline testing.
These documents are uploaded through the UI AFTER the TM has been seeded.
Run AFTER seed_translation_memory.py.

Creates:
  - 04_software_guide_v2.docx   — Paraphrased software guide (fuzzy TM matches)
  - 05_hr_policy_v2.docx        — Paraphrased HR policy (fuzzy TM matches + validation issues)
  - 06_new_product_launch.docx  — Completely new content (no TM matches, full LLM translation)

Usage:
    python test_data/create_test_documents.py
"""

import os
from docx import Document


DOCS_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "docs")


def ensure_docs_dir():
    os.makedirs(DOCS_DIR, exist_ok=True)
    print(f"[OK] Docs directory ready: {DOCS_DIR}")


def create_software_guide_v2():
    """
    Create 04_software_guide_v2.docx — ~300 words.
    Paraphrased version of doc 01 to trigger fuzzy TM matches.

    Deliberate issues:
      - Two instances of double spaces
      - "application" and "app" used interchangeably (terminology variation)
    """
    doc = Document()

    # --- Introduction ---
    doc.add_heading("Introduction", level=1)
    #                                        ↓↓ deliberate double space
    doc.add_paragraph(
        "CloudSync is a robust cloud-based  file management and collaboration tool built for modern teams. "
        "It allows users to share files instantly, collaborate in real time, and maintain automatic backups "
        "across all connected devices."
    )

    # --- System Requirements ---
    doc.add_heading("System Requirements", level=1)
    doc.add_paragraph(
        "Please verify that your computer meets these requirements before installing the app:"
    )
    doc.add_paragraph("Operating System: Windows 10+, macOS 12+, or Ubuntu 20.04+", style="List Bullet")
    #                                                         ↓ paraphrased from "requires a minimum of 4GB RAM to operate"
    doc.add_paragraph("RAM: CloudSync needs at least 4GB of RAM to run properly. We recommend 8GB for best performance.", style="List Bullet")
    doc.add_paragraph("Storage: You need a minimum of 500MB of free disk space for the application itself.", style="List Bullet")
    doc.add_paragraph("Internet: A stable broadband connection of at least 10 Mbps is necessary for syncing.", style="List Bullet")

    # --- Installation ---
    doc.add_heading("Installation", level=1)
    #                                                                                    ↓↓ deliberate double space
    doc.add_paragraph(
        "Download the latest version of the application from the official CloudSync website.  "
        "Run the installer and follow the prompts to complete the setup. "
        # ↓ "app" instead of "application" — terminology variation
        "After installation, restart the app to apply all changes. "
        "The system automatically saves your installation progress if the process gets interrupted."
    )

    # --- Getting Started ---
    doc.add_heading("Getting Started", level=1)
    doc.add_paragraph(
        "Once installed, open CloudSync and register a new account with your work email. "
        "You will receive a confirmation email — your account is ready once you verify it. "
        # ↓ "app" instead of "application" — terminology variation
        "Go to Settings in the app to set up your notification preferences, sync schedule, "
        "and default storage folders."
    )

    # --- Troubleshooting ---
    doc.add_heading("Troubleshooting", level=1)
    doc.add_paragraph(
        "If you run into problems, check the troubleshooting section below. "
        "Restarting the application or verifying your network connection resolves most issues. "
        "Reach out to our support team if the issue continues after trying these steps."
    )

    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    table.cell(0, 0).text = "Error Code"
    table.cell(0, 1).text = "Description"
    table.cell(1, 0).text = "ERR-101"
    table.cell(1, 1).text = "Connection timed out. Verify your internet and firewall configuration."
    table.cell(2, 0).text = "ERR-205"
    table.cell(2, 1).text = "Login failed. Double-check your credentials and retry."
    table.cell(3, 0).text = "ERR-310"
    table.cell(3, 1).text = "Not enough storage. Clear disk space or ask your admin for help."

    filepath = os.path.join(DOCS_DIR, "04_software_guide_v2.docx")
    doc.save(filepath)
    print(f"[OK] Created: {filepath}")
    print("     → Fuzzy TM matches expected (paraphrased software domain)")
    print("     → Deliberate issues: 2x double spaces, 'application'/'app' interchangeably")
    return filepath


def create_hr_policy_v2():
    """
    Create 05_hr_policy_v2.docx — ~300 words.
    Paraphrased version of doc 02 to trigger fuzzy TM matches.

    Deliberate issues:
      - Mixed date formats: DD/MM/YYYY and Month DD YYYY
      - "Employee" and "Staff Member" used interchangeably
    """
    doc = Document()

    # --- Code of Conduct ---
    doc.add_heading("Employee Code of Conduct", level=1)
    doc.add_paragraph(
        #                                                   ↓ paraphrased from "zero-tolerance policy"
        "Our organization enforces a strict zero-tolerance approach to harassment, discrimination, and any "
        "form of unethical conduct. All staff members are expected to behave with professionalism and respect "
        #         ↓ "Staff members" interchangeably with "Employees"
        "toward colleagues and clients. Staff members must follow all applicable health and safety rules."
    )

    # --- Leave Policy ---
    doc.add_heading("Leave Policy", level=1)
    doc.add_paragraph(
        #                                 ↓ paraphrased from "entitled to 20 days"
        "Each employee is eligible for 20 working days of annual leave every calendar year. "
        #                                                            ↓ DD/MM/YYYY format (deliberate)
        "Leave applications must be filed at least 14 days before the start date via the HR portal. "
        "The deadline for carrying forward unused leave for 2025 is 31/03/2026."
    )
    doc.add_paragraph("Annual Leave — 20 working days per year", style="List Bullet")
    doc.add_paragraph("Sick Leave — 10 days per year (medical certificate required)", style="List Bullet")
    doc.add_paragraph("Maternity Leave — 26 weeks as per legal requirements", style="List Bullet")
    doc.add_paragraph("Paternity Leave — 2 weeks following the birth or adoption of a child", style="List Bullet")
    doc.add_paragraph("Emergency Leave — Up to 5 days for personal emergencies", style="List Bullet")

    # --- Performance Review ---
    doc.add_heading("Performance Review Process", level=1)
    doc.add_paragraph(
        #                                                        ↓ paraphrased from "conducted on a quarterly basis"
        "Employee performance reviews take place every three months to track progress against targets. "
        "Each review consists of a self-assessment, manager feedback, and a goal-setting discussion. "
        #                  ↓ "Staff Member" instead of "Employee" — terminology variation
        #                                                                  ↓ Month DD YYYY format (deliberate)
        "The next review cycle for all staff members begins on January 15 2026."
    )

    # --- Grievance Procedure ---
    doc.add_heading("Grievance Procedure", level=1)
    doc.add_paragraph(
        #                                 ↓ paraphrased from "submitted in writing to HR"
        "All workplace complaints must be filed in writing through the official grievance form on the intranet. "
        "HR will confirm receipt within three working days and begin the investigation process. "
        #                                                      ↓ DD/MM/YYYY format (deliberate)
        "All grievances submitted before 15/06/2026 will be processed under the current policy framework."
    )

    table = doc.add_table(rows=4, cols=3)
    table.style = "Table Grid"
    table.cell(0, 0).text = "Review Type"
    table.cell(0, 1).text = "Frequency"
    table.cell(0, 2).text = "Responsible Party"
    table.cell(1, 0).text = "Quarterly Review"
    table.cell(1, 1).text = "Every 3 months"
    table.cell(1, 2).text = "Direct Manager"
    table.cell(2, 0).text = "Annual Review"
    table.cell(2, 1).text = "Once per year"
    table.cell(2, 2).text = "HR Department"
    table.cell(3, 0).text = "Probation Review"
    table.cell(3, 1).text = "End of probation"
    table.cell(3, 2).text = "Manager & HR"

    filepath = os.path.join(DOCS_DIR, "05_hr_policy_v2.docx")
    doc.save(filepath)
    print(f"[OK] Created: {filepath}")
    print("     → Fuzzy TM matches expected (paraphrased HR domain)")
    print("     → Deliberate issues: mixed date formats (DD/MM/YYYY + Month DD YYYY), 'Employee'/'Staff Member' interchangeably")
    return filepath


def create_new_product_launch():
    """
    Create 06_new_product_launch.docx — ~300 words.
    Completely new content with NO matching TM entries.
    Tests the full LLM translation path.
    """
    doc = Document()

    # --- Product Launch Announcement ---
    doc.add_heading("Product Launch Announcement", level=1)
    doc.add_paragraph(
        "We are thrilled to announce the launch of NovaPay, our next-generation payment processing platform "
        "designed to simplify financial transactions for businesses of all sizes. NovaPay combines cutting-edge "
        "security technology with an intuitive user interface to deliver a seamless payment experience. "
        "After two years of development and extensive beta testing with over 500 businesses, NovaPay is now "
        "ready for general availability."
    )

    # --- Key Features ---
    doc.add_heading("Key Features", level=1)
    doc.add_paragraph("Instant Payment Processing — Transactions are settled in under 3 seconds with real-time confirmation notifications.", style="List Bullet")
    doc.add_paragraph("Multi-Currency Support — Accept and process payments in over 150 currencies with automatic exchange rate conversion.", style="List Bullet")
    doc.add_paragraph("Advanced Fraud Detection — AI-powered fraud monitoring analyzes transaction patterns and flags suspicious activity in real time.", style="List Bullet")
    doc.add_paragraph("Developer-Friendly API — Comprehensive RESTful API with SDKs for Python, JavaScript, Ruby, and Java for seamless integration.", style="List Bullet")
    doc.add_paragraph("Compliance Dashboard — Built-in compliance tools ensure adherence to PCI DSS, GDPR, and PSD2 regulations across all markets.", style="List Bullet")

    # --- Pricing ---
    doc.add_heading("Pricing", level=1)
    doc.add_paragraph(
        "NovaPay offers flexible pricing plans to accommodate businesses at every growth stage:"
    )

    table = doc.add_table(rows=4, cols=3)
    table.style = "Table Grid"
    table.cell(0, 0).text = "Plan"
    table.cell(0, 1).text = "Price"
    table.cell(0, 2).text = "Features Included"
    table.cell(1, 0).text = "Starter"
    table.cell(1, 1).text = "$29/month"
    table.cell(1, 2).text = "Up to 1,000 transactions, basic fraud detection, email support"
    table.cell(2, 0).text = "Business"
    table.cell(2, 1).text = "$99/month"
    table.cell(2, 2).text = "Up to 10,000 transactions, advanced fraud detection, priority support, multi-currency"
    table.cell(3, 0).text = "Enterprise"
    table.cell(3, 1).text = "Custom pricing"
    table.cell(3, 2).text = "Unlimited transactions, dedicated account manager, SLA guarantee, custom integrations"

    # --- Launch Timeline ---
    doc.add_heading("Launch Timeline", level=1)
    doc.add_paragraph(
        "NovaPay will be available in three phases to ensure a smooth global rollout. "
        "Phase 1 covers North America and Europe, launching on May 1, 2026. "
        "Phase 2 will expand to Asia-Pacific and Latin America markets by July 15, 2026, with full "
        "global availability targeted for September 30, 2026."
    )

    # --- Contact ---
    doc.add_heading("Contact", level=1)
    doc.add_paragraph(
        "For more information about NovaPay, please visit our website at www.novapay.example.com "
        "or email our sales team at sales@novapay.example.com. "
        "Enterprise customers can schedule a personalized demo by contacting enterprise@novapay.example.com."
    )

    filepath = os.path.join(DOCS_DIR, "06_new_product_launch.docx")
    doc.save(filepath)
    print(f"[OK] Created: {filepath}")
    print("     → NO TM matches expected (completely new content)")
    print("     → All segments go to LLM for full translation")
    return filepath


if __name__ == "__main__":
    print("=" * 60)
    print("STEP 3: Creating Test Documents for Pipeline Testing")
    print("=" * 60)

    ensure_docs_dir()
    create_software_guide_v2()
    create_hr_policy_v2()
    create_new_product_launch()

    print()
    print("=" * 60)
    print("All 3 test documents created successfully!")
    print("=" * 60)
    print()
    print("NEXT: Upload these documents through the UI:")
    print("  - 04_software_guide_v2.docx   → Tests fuzzy TM matching (software domain)")
    print("  - 05_hr_policy_v2.docx        → Tests fuzzy TM matching + validation issues (HR domain)")
    print("  - 06_new_product_launch.docx  → Tests full LLM translation (no TM matches)")
    print()
    print("=" * 60)
