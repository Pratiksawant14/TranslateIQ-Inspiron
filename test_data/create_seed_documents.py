"""
STEP 1: Create seed documents for building Translation Memory.
These 3 documents provide the reference corpus that TM pairs are based on.
Run this FIRST before seed_translation_memory.py.

Usage:
    python test_data/create_seed_documents.py
"""

import os
from docx import Document


DOCS_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "docs")


def ensure_docs_dir():
    os.makedirs(DOCS_DIR, exist_ok=True)
    print(f"[OK] Docs directory ready: {DOCS_DIR}")


def create_software_user_guide():
    """Create 01_software_user_guide.docx — ~400 words, software domain."""
    doc = Document()

    # --- Introduction ---
    doc.add_heading("Introduction", level=1)
    doc.add_paragraph(
        "CloudSync is a powerful cloud-based file synchronization and collaboration platform "
        "designed for teams of all sizes. It enables seamless file sharing, real-time collaboration, "
        "and automated backup across multiple devices. With CloudSync, your team can work together "
        "more efficiently while keeping all data secure and accessible."
    )

    # --- System Requirements ---
    doc.add_heading("System Requirements", level=1)
    doc.add_paragraph(
        "Before installing CloudSync, please ensure your system meets the following minimum requirements:"
    )
    doc.add_paragraph("Operating System: Windows 10 or later, macOS 12 or later, Ubuntu 20.04 or later", style="List Bullet")
    doc.add_paragraph("RAM: CloudSync requires a minimum of 4GB RAM to operate. 8GB RAM is recommended for optimal performance.", style="List Bullet")
    doc.add_paragraph("Storage: At least 500MB of free disk space is required for the application. Additional space is needed for synchronized files.", style="List Bullet")
    doc.add_paragraph("Internet Connection: A stable broadband internet connection with a minimum speed of 10 Mbps is required for real-time synchronization.", style="List Bullet")

    # --- Installation ---
    doc.add_heading("Installation", level=1)
    doc.add_paragraph(
        "To install CloudSync, download the latest installer from the official website at www.cloudsync.example.com. "
        "Run the downloaded installer and follow the on-screen instructions to complete the setup process. "
        "Please restart the application to apply the changes after the installation is complete. "
        "The system will automatically save your progress during the installation if the process is interrupted."
    )

    # --- Getting Started ---
    doc.add_heading("Getting Started", level=1)
    doc.add_paragraph(
        "After installation, launch CloudSync and create a new account using your corporate email address. "
        "Your account has been successfully created once you verify the confirmation email sent to your inbox. "
        "Click the Settings icon to configure your preferences, including notification settings, sync frequency, "
        "and default file storage locations."
    )

    # --- Troubleshooting ---
    doc.add_heading("Troubleshooting", level=1)
    doc.add_paragraph(
        "If you encounter any issues during installation or usage, please consult the troubleshooting guide below. "
        "Most common issues can be resolved by restarting the application or checking your internet connection. "
        "Contact our support team for further assistance if the problem persists after following the steps below."
    )

    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    # Header row
    table.cell(0, 0).text = "Error Code"
    table.cell(0, 1).text = "Description"
    # Data rows
    table.cell(1, 0).text = "ERR-101"
    table.cell(1, 1).text = "Connection timeout. Check your internet connection and firewall settings."
    table.cell(2, 0).text = "ERR-205"
    table.cell(2, 1).text = "Authentication failed. Verify your login credentials and try again."
    table.cell(3, 0).text = "ERR-310"
    table.cell(3, 1).text = "Insufficient storage space. Free up disk space or contact your administrator."

    filepath = os.path.join(DOCS_DIR, "01_software_user_guide.docx")
    doc.save(filepath)
    print(f"[OK] Created: {filepath}")
    return filepath


def create_hr_policy_document():
    """Create 02_hr_policy_document.docx — ~400 words, HR domain."""
    doc = Document()

    # --- Employee Code of Conduct ---
    doc.add_heading("Employee Code of Conduct", level=1)
    doc.add_paragraph(
        "The company maintains a strict zero-tolerance policy regarding harassment, discrimination, and unethical behavior "
        "in the workplace. All employees are expected to conduct themselves with professionalism, integrity, and respect "
        "toward their colleagues, clients, and business partners. Employees must comply with all workplace health and safety "
        "regulations as outlined in this document and any applicable local laws."
    )

    # --- Leave Policy ---
    doc.add_heading("Leave Policy", level=1)
    doc.add_paragraph(
        "Employees are entitled to 20 days of annual leave per year, accrued on a monthly basis. "
        "Leave requests must be submitted at least two weeks in advance through the company's HR management portal. "
        "Unused annual leave may be carried forward to the following year, up to a maximum of five additional days."
    )
    doc.add_paragraph("Annual Leave — 20 working days per calendar year", style="List Bullet")
    doc.add_paragraph("Sick Leave — 10 working days per calendar year with valid medical certification", style="List Bullet")
    doc.add_paragraph("Maternity Leave — 26 weeks of paid leave as per statutory requirements", style="List Bullet")
    doc.add_paragraph("Paternity Leave — 2 weeks of paid leave following the birth or adoption of a child", style="List Bullet")
    doc.add_paragraph("Emergency Leave — Up to 5 days per year for unforeseen personal emergencies", style="List Bullet")

    # --- Performance Review Process ---
    doc.add_heading("Performance Review Process", level=1)
    doc.add_paragraph(
        "Performance reviews are conducted on a quarterly basis to assess employee progress against established goals. "
        "Each review includes a self-assessment component, a manager evaluation, and a collaborative goal-setting session. "
        "All grievances related to performance evaluations must be submitted in writing to HR within 10 business days "
        "of the review meeting."
    )

    # --- Grievance Procedure ---
    doc.add_heading("Grievance Procedure", level=1)
    doc.add_paragraph(
        "All grievances must be submitted in writing to HR using the official grievance form available on the company intranet. "
        "The HR department will acknowledge receipt of the grievance within three business days and initiate an investigation. "
        "A formal response will be provided within 15 business days of the initial submission, and the employee will be "
        "informed of any further actions or resolutions."
    )

    table = doc.add_table(rows=4, cols=3)
    table.style = "Table Grid"
    # Header row
    table.cell(0, 0).text = "Review Type"
    table.cell(0, 1).text = "Frequency"
    table.cell(0, 2).text = "Responsible Party"
    # Data rows
    table.cell(1, 0).text = "Quarterly Performance Review"
    table.cell(1, 1).text = "Every 3 months"
    table.cell(1, 2).text = "Direct Manager"
    table.cell(2, 0).text = "Annual Comprehensive Review"
    table.cell(2, 1).text = "Once per year"
    table.cell(2, 2).text = "HR Department"
    table.cell(3, 0).text = "Probation Review"
    table.cell(3, 1).text = "End of probation period"
    table.cell(3, 2).text = "Direct Manager & HR"

    filepath = os.path.join(DOCS_DIR, "02_hr_policy_document.docx")
    doc.save(filepath)
    print(f"[OK] Created: {filepath}")
    return filepath


def create_legal_terms():
    """Create 03_legal_terms.docx — ~400 words, legal domain."""
    doc = Document()

    # --- Terms and Conditions ---
    doc.add_heading("Terms and Conditions", level=1)
    doc.add_paragraph(
        "This agreement is governed by applicable law and establishes the terms and conditions under which services "
        "will be provided by the Company to the Client. By signing this document, all parties acknowledge that they "
        "have read, understood, and agreed to the provisions contained herein."
    )

    # --- Payment Terms ---
    doc.add_heading("Payment Terms", level=1)
    doc.add_paragraph(
        "Payment shall be made within thirty days of invoice in accordance with the schedule outlined below. "
        "All payments must be made in the agreed-upon currency via wire transfer or other approved payment methods. "
        "Late payments will incur a penalty of 1.5% per month on the outstanding balance, calculated from the due date."
    )

    table_pay = doc.add_table(rows=4, cols=3)
    table_pay.style = "Table Grid"
    table_pay.cell(0, 0).text = "Milestone"
    table_pay.cell(0, 1).text = "Amount"
    table_pay.cell(0, 2).text = "Due Date"
    table_pay.cell(1, 0).text = "Project Initiation"
    table_pay.cell(1, 1).text = "$25,000"
    table_pay.cell(1, 2).text = "Upon contract signing"
    table_pay.cell(2, 0).text = "Mid-Project Delivery"
    table_pay.cell(2, 1).text = "$50,000"
    table_pay.cell(2, 2).text = "90 days after initiation"
    table_pay.cell(3, 0).text = "Final Delivery"
    table_pay.cell(3, 1).text = "$25,000"
    table_pay.cell(3, 2).text = "Upon project completion"

    # --- Termination Clause ---
    doc.add_heading("Termination Clause", level=1)
    doc.add_paragraph(
        "Either party may terminate this agreement with written notice of at least thirty calendar days. "
        "In the event of termination, all work completed up to the date of termination shall be compensated "
        "in accordance with the payment schedule. The terminating party shall not be liable for any consequential "
        "damages arising from the termination of this agreement."
    )

    # --- Governing Law ---
    doc.add_heading("Governing Law", level=1)
    doc.add_paragraph(
        "This agreement is governed by applicable law of the State of California, United States of America. "
        "The parties agree to resolve disputes through arbitration in San Francisco, California, in accordance "
        "with the rules of the American Arbitration Association."
    )

    # --- Confidentiality ---
    doc.add_heading("Confidentiality", level=1)
    doc.add_paragraph(
        "All confidential information must be protected from disclosure to unauthorized third parties for the "
        "duration of this agreement and for a period of five years following its termination. "
        "Confidential information includes, but is not limited to, trade secrets, business strategies, financial data, "
        "and proprietary technology. Any breach of this confidentiality clause may result in immediate termination "
        "of the agreement and legal action for damages."
    )

    filepath = os.path.join(DOCS_DIR, "03_legal_terms.docx")
    doc.save(filepath)
    print(f"[OK] Created: {filepath}")
    return filepath


if __name__ == "__main__":
    print("=" * 60)
    print("STEP 1: Creating Seed Documents for Translation Memory")
    print("=" * 60)

    ensure_docs_dir()
    create_software_user_guide()
    create_hr_policy_document()
    create_legal_terms()

    print()
    print("=" * 60)
    print("All 3 seed documents created successfully!")
    print("Next step: Run seed_translation_memory.py")
    print("=" * 60)
